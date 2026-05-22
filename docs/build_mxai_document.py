from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
DOCX_PATH = OUT_DIR / "MXAI_Detector_Documentacion_Tecnica.docx"


COLORS = {
    "cyan": "28D8D1",
    "dark": "0B1720",
    "ink": "17232D",
    "muted": "5E6F78",
    "line": "D7E1E5",
    "fill": "EFF7F8",
    "soft": "F4F6F8",
    "danger": "A23B52",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D7E1E5", size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("MXAI Detector - Documentacion tecnica")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x5E, 0x6F, 0x78)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_code(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = False
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0B1720")
    set_cell_border(cell, "24404A", "8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xE7, 0xF2, 0xF2)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = "EFF7F8"):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x37, 0x45)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [2700, 6660])
    hdr = table.rows[0].cells
    hdr[0].text = "Elemento"
    hdr[1].text = "Descripcion"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for label, desc in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = desc
    set_table_width(table, [2700, 6660])
    doc.add_paragraph()


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("MXAI Detector")
    r.font.name = "Calibri"
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run(
        "Documentacion tecnica del programa mobile de analisis Android con XAI"
    )
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x5E, 0x6F, 0x78)

    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, [2300, 7060])
    for row, values in enumerate(
        [
            ("Proyecto", "Mobile X AI Detector"),
            ("Plataforma", "Flutter + Android nativo + Ollama + backend Python opcional"),
            ("Modelo ML", "Android_Malwaredetector/android_malware_detector.joblib"),
            ("Fecha", "22 de mayo de 2026"),
        ]
    ):
        meta.cell(row, 0).text = values[0]
        meta.cell(row, 1).text = values[1]
        set_cell_shading(meta.cell(row, 0), "E8EEF5")
    doc.add_paragraph()

    add_callout(
        doc,
        "Resumen",
        "MXAI Detector es una app Flutter orientada a analisis de comportamiento Android. "
        "La interfaz ejecuta un analisis global del dispositivo, recibe APKs externos por "
        "Android intents, muestra una decision Benign/Malware y usa Ollama para transformar "
        "probabilidades y factores locales en explicaciones legibles.",
    )

    add_heading(doc, "1. Proposito del programa", 1)
    add_body(
        doc,
        "El programa busca combinar deteccion de malware Android con explicabilidad. "
        "El usuario no selecciona manualmente una app desde la pantalla principal: pulsa "
        "Analizar y el sistema observa apps, permisos y senales del dispositivo para producir "
        "un resultado global. Para APKs descargados, el flujo se activa fuera de la app mediante "
        "Compartir/Enviar a MXAI despues de la descarga.",
    )
    add_body(
        doc,
        "La meta funcional final es conectar un extractor de features compatible con el "
        "modelo entrenado, producir un vector de 470 columnas y ejecutar el modelo "
        "android_malware_detector.joblib. En el estado actual, Flutter usa un motor prototipo "
        "calibrado y deja preparada una API Python para inferencia real con joblib.",
    )

    add_heading(doc, "2. Flujo conceptual de analisis", 1)
    add_code(
        doc,
        "APK o comportamiento observado\n"
        "  -> Extractor de features compatible\n"
        "  -> Vector de 470 columnas\n"
        "  -> android_malware_detector.joblib\n"
        "  -> Prediccion Benign/Malware\n"
        "  -> SHAP / importancia local\n"
        "  -> Ollama explica el resultado",
    )
    add_body(
        doc,
        "El punto central del flujo es la compatibilidad con las 470 columnas del modelo. "
        "La capa visual no decide por si sola: presenta la salida del motor de deteccion y "
        "separa la explicacion en dos funciones: reporte XAI del analisis global y chatbot "
        "defensivo de ciberseguridad.",
    )

    add_heading(doc, "3. Arquitectura general", 1)
    add_kv_table(
        doc,
        [
            (
                "Flutter UI",
                "Pantallas, scanner circular, navbar lateral, reporte XAI, chatbot y ventana externa de APK.",
            ),
            (
                "Android nativo",
                "MethodChannel para listar apps instaladas, abrir ajustes de uso y recibir APKs compartidos.",
            ),
            (
                "Motor prototipo",
                "PrototypeDetectionEngine genera una decision temporal mientras se conecta el extractor real.",
            ),
            (
                "Modelo ML",
                "Archivos de metadata, features e importancia global se cargan como assets; el joblib queda para backend nativo/Python.",
            ),
            (
                "Ollama",
                "mxai-xai-report genera reportes; mxai-cyber-chat responde preguntas defensivas; mxai-tinyllama-local valida el GGUF.",
            ),
            (
                "Backend Python",
                "FastAPI recibe vector o diccionario de features, ejecuta joblib y devuelve probabilidades e importancia local.",
            ),
        ],
    )

    add_heading(doc, "4. Estructura de archivos principal", 1)
    add_kv_table(
        doc,
        [
            ("lib/main.dart", "Interfaz principal, pantallas, navegacion lateral y paneles de analisis."),
            ("lib/models/detector_models.dart", "Modelos de datos: metadata, target, resultado y factores locales."),
            ("lib/services/detector_service.dart", "Motor temporal de deteccion y calibracion conservadora."),
            ("lib/services/android_app_scanner_service.dart", "Cliente Flutter del MethodChannel Android."),
            ("lib/services/ollama_service.dart", "Cliente HTTP hacia /api/generate de Ollama."),
            ("backend/inference_api.py", "Servidor FastAPI para ejecutar el modelo joblib con un vector compatible."),
            ("Android_Malwaredetector/", "Modelo entrenado, features, metricas, mapping de clases e importancias."),
            ("ollama/", "Modelfiles para reporte, chatbot y validacion del GGUF local."),
            ("scripts/", "Automatizacion para crear modelos Ollama e instalar el APK por ADB."),
        ],
    )

    add_heading(doc, "5. Interfaz movil", 1)
    add_body(
        doc,
        "La app usa una estetica futurista oscura, con fondo plano de baja interferencia, "
        "scanner circular y paneles compactos. La navbar vertical se abre al tocar el escudo "
        "y desplaza ligeramente el contenido para mantener orientacion espacial.",
    )
    add_bullet(doc, "Inicio: analisis global, scanner circular, boton Analizar y reporte XAI.")
    add_bullet(doc, "Decision: explica como decide el sistema y muestra el resultado actual.")
    add_bullet(doc, "Datos: resume senales observadas, features y factores locales.")
    add_bullet(doc, "Chat XAI: chatbot limitado a ciberseguridad defensiva y preguntas sobre MXAI.")
    add_bullet(doc, "SDK Android: permisos, servicios y flujo externo para APK descargado.")

    add_heading(doc, "6. Analisis global", 1)
    add_body(
        doc,
        "El analisis global evita falsos positivos por volumen bruto de permisos. Android y "
        "los fabricantes preinstalan muchas apps con permisos sensibles; por eso el motor "
        "prototipo separa apps de usuario y apps de sistema/OEM, mide concentracion de senales "
        "y exige un umbral mas alto antes de marcar Malware.",
    )
    add_bullet(doc, "Si no hay senales fuertes, el riesgo global queda en rango Benign.")
    add_bullet(doc, "Si varias apps de usuario concentran permisos y senales de alto riesgo, el resultado puede elevarse.")
    add_bullet(doc, "El reporte XAI aparece en Inicio y usa solo los datos del analisis generado.")

    add_heading(doc, "7. Analisis de APKs descargados", 1)
    add_body(
        doc,
        "MXAI ya no compite con la descarga ni con el instalador del sistema. El manifest no "
        "declara ACTION_VIEW para APK; declara ACTION_SEND. Eso fuerza el flujo correcto: "
        "primero el usuario descarga el APK, despues lo comparte o envia a MXAI para analizarlo.",
    )
    add_code(doc, "Descarga del APK\n  -> Compartir / Enviar a MXAI\n  -> Ventana externa de analisis\n  -> Decision + explicabilidad")
    add_body(
        doc,
        "La ventana externa reutiliza la experiencia visual del scanner, pero trabaja con el "
        "archivo APK recibido por Android intent. La app valida nombre, URI, tamano y MIME para "
        "confirmar que el archivo corresponde a un APK.",
    )

    add_heading(doc, "8. Modelo ML y backend de inferencia", 1)
    add_body(
        doc,
        "El archivo android_malware_detector.joblib es un modelo scikit-learn serializado. "
        "Dart/Flutter no puede ejecutarlo directamente de forma nativa. Por eso el proyecto "
        "mantiene dos caminos: un motor prototipo para la interfaz movil y una API Python para "
        "ejecutar inferencia real cuando exista el vector de 470 features.",
    )
    add_code(
        doc,
        "python -m pip install -r backend\\requirements.txt\n"
        "python -m uvicorn backend.inference_api:app --reload --host 127.0.0.1 --port 8000",
    )
    add_body(
        doc,
        "El endpoint POST /predict acepta vector con 470 valores o un diccionario de features. "
        "Devuelve prediccion, probabilidades Benign/Malware e importancia local. Si SHAP esta "
        "disponible, usa TreeExplainer; si no, calcula una aproximacion con feature_importances_.",
    )

    add_heading(doc, "9. Ollama y XAI", 1)
    add_body(
        doc,
        "La app llama a Ollama mediante HTTP en /api/generate. En telefono fisico, 127.0.0.1 "
        "apunta al telefono, no a la PC. Por eso el APK se compila con OLLAMA_BASE_URL apuntando "
        "a la IP local del equipo donde corre Ollama.",
    )
    add_kv_table(
        doc,
        [
            ("mxai-xai-report", "Genera reporte del analisis global usando prediccion, probabilidades y factores locales."),
            ("mxai-cyber-chat", "Responde preguntas defensivas de ciberseguridad y sobre la app MXAI."),
            ("mxai-tinyllama-local", "Carga tinyllama.Q4_K_M.gguf como validacion de archivo local; no es el modelo activo si su salida es corrupta."),
        ],
    )
    add_code(
        doc,
        ".\\scripts\\create_ollama_models.ps1\n"
        "$env:OLLAMA_HOST=\"0.0.0.0:11434\"\n"
        "ollama serve\n"
        ".\\scripts\\install_to_phone.ps1 -OllamaBaseUrl \"http://TU_IP_LOCAL:11434\"",
    )
    add_callout(
        doc,
        "Nota practica sobre Ollama",
        "El APK no instala Ollama dentro de Android. Ollama corre como servidor externo en la PC o en otro host. "
        "Para una version 100% offline dentro del telefono, el siguiente paso seria integrar un runtime nativo "
        "como llama.cpp, ONNX Runtime o TFLite y cargar el modelo como asset compatible.",
        "FFF7E6",
    )

    add_heading(doc, "10. Permisos y servicios Android", 1)
    add_kv_table(
        doc,
        [
            ("INTERNET", "Permite conectar con Ollama o con un backend de inferencia en la red local."),
            ("QUERY_ALL_PACKAGES", "Permite listar apps instaladas durante pruebas locales."),
            ("PACKAGE_USAGE_STATS", "Prepara acceso a estadisticas de uso; requiere activacion manual por el usuario."),
            ("ACTION_SEND APK", "Permite recibir APKs despues de descargarlos y compartirlos con MXAI."),
            ("AppAnalysisService", "Servicio reservado para analisis en segundo plano cuando se conecte el extractor real."),
        ],
    )

    add_heading(doc, "11. Instalacion y prueba", 1)
    add_number(doc, "Conectar el telefono por USB y activar depuracion USB.")
    add_number(doc, "Crear los modelos Ollama con scripts\\create_ollama_models.ps1.")
    add_number(doc, "Iniciar Ollama escuchando en red local con OLLAMA_HOST=0.0.0.0:11434.")
    add_number(doc, "Instalar con scripts\\install_to_phone.ps1; el script detecta la IP local si no se pasa una manual.")
    add_number(doc, "Abrir la app, ejecutar Analizar y despues generar el reporte XAI.")
    add_code(
        doc,
        "flutter analyze\nflutter test\nflutter build web\nflutter build apk --debug",
    )

    add_heading(doc, "12. Estado actual y limitaciones", 1)
    add_bullet(doc, "La interfaz movil, navbar lateral, scanner, reporte XAI y chatbot ya estan implementados.")
    add_bullet(doc, "El motor de Flutter sigue siendo prototipo hasta conectar el extractor de 470 columnas.")
    add_bullet(doc, "El joblib se ejecuta desde backend Python o debera convertirse a ONNX/TFLite/plugin nativo.")
    add_bullet(doc, "El archivo tinyllama.Q4_K_M.gguf se importa en Ollama, pero debe reemplazarse por una variante instruct/chat si devuelve salida corrupta.")
    add_bullet(doc, "La reduccion de falsos positivos en analisis global se basa en calibracion conservadora y separacion user/system apps.")

    add_heading(doc, "13. Roadmap recomendado", 1)
    add_number(doc, "Construir extractor real que emita exactamente las 470 features esperadas.")
    add_number(doc, "Conectar Flutter al backend /predict o portar el modelo a un formato movil.")
    add_number(doc, "Agregar SHAP local real en la respuesta de inferencia y mapearlo al panel de Datos.")
    add_number(doc, "Empaquetar una opcion offline de XAI con runtime nativo si se requiere funcionamiento sin PC.")
    add_number(doc, "Endurecer permisos, privacidad y comunicacion de red antes de distribucion fuera de laboratorio.")

    add_heading(doc, "14. Glosario breve", 1)
    add_kv_table(
        doc,
        [
            ("APK", "Paquete instalable de Android."),
            ("Feature", "Variable numerica o categorica usada por el modelo."),
            ("Vector 470", "Entrada ordenada que el modelo entrenado espera recibir."),
            ("SHAP", "Tecnica de explicabilidad que estima contribuciones locales de features."),
            ("LLM", "Modelo de lenguaje usado para redactar explicaciones a partir de datos tecnicos."),
            ("Ollama", "Servidor local que ejecuta modelos de lenguaje y expone /api/generate."),
        ],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
    print(DOCX_PATH)
